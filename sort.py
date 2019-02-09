import re
from collections import defaultdict, OrderedDict
from docx import Document

''' Open the word document '''

document = Document('references.docx')
references = defaultdict(list) # Using lists to hold values under the same year as the key.

''' For every paragraph in the document try to find any four digit number
	where the first digit should either be 1 or 2 and the rest of the 
	digits are 0 - 9, and add it to the references dictionary with the 
	matched string as the key.
''' 
for line in document.paragraphs:
	line = line.text.encode('ascii', 'ignore').decode('ascii')
	match = re.search(r'([1-2][0-9]{3})', line)
	
	if match:
		year = match.group(1)
	else:
		year = "0000"
	
	references[year].append(line)

''' Sort by keys '''
ordered_references = OrderedDict(sorted(references.items(), reverse=True))

''' Write the sorted values to a new document. '''
new_document = Document()

for year, references in ordered_references.iteritems():
	for reference in references:
		new_document.add_paragraph(reference)

new_document.save('sorted_references.docx')